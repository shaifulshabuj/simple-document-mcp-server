#!/usr/bin/env python3
"""
Simple Document MCP Server
A minimal MCP (Model Context Protocol) server for document processing and search.
"""

import asyncio
import json
import logging
import argparse
import sys
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
import mcp.server.stdio
import mcp.types as types
from mcp.server import NotificationOptions, Server
from mcp.server.models import InitializationOptions

# Document processing imports
try:
    import PyPDF2
    import docx
    import openpyxl
    from langdetect import detect
    import re
except ImportError as e:
    print(f"Missing required dependency: {e}")
    print("Please install dependencies with: pip install -r requirements.txt")
    exit(1)

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentInfo:
    """Document information structure"""
    path: str
    filename: str
    content: str
    language: str
    file_type: str
    size: int
    last_modified: float

class SimpleDocumentProcessor:
    """Document processor for multiple formats"""
    
    def __init__(self, documents_dir: str = "./documents"):
        self.documents_dir = Path(documents_dir)
        self.documents_dir.mkdir(exist_ok=True)
        self.document_cache: Dict[str, DocumentInfo] = {}
        
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def extract_text_from_xlsx(self, file_path: Path) -> str:
        """Extract text from XLSX file"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text = ""
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading XLSX {file_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file with multiple encoding support"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'shift_jis', 'cp932']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading TXT {file_path} with {encoding}: {e}")
                continue
        
        logger.error(f"Could not read {file_path} with any encoding")
        return ""
    
    def detect_language(self, text: str) -> str:
        """Detect language of text"""
        try:
            # Clean text for better detection
            clean_text = re.sub(r'[^\w\s]', '', text)
            if len(clean_text.strip()) > 20:
                detected = detect(clean_text)
                return detected
            return "unknown"
        except Exception:
            return "unknown"
    
    def process_document(self, file_path: Path) -> Optional[DocumentInfo]:
        """Process a single document"""
        if not file_path.exists():
            return None
        
        file_extension = file_path.suffix.lower()
        
        # Extract text based on file type
        extractors = {
            '.pdf': (self.extract_text_from_pdf, "PDF"),
            '.docx': (self.extract_text_from_docx, "Word Document"),
            '.xlsx': (self.extract_text_from_xlsx, "Excel Spreadsheet"),
            '.txt': (self.extract_text_from_txt, "Text File"),
        }
        
        if file_extension not in extractors:
            logger.warning(f"Unsupported file type: {file_extension}")
            return None
        
        extractor_func, file_type = extractors[file_extension]
        text = extractor_func(file_path)
        
        if not text:
            logger.warning(f"No text extracted from {file_path}")
            return None
        
        # Detect language
        language = self.detect_language(text)
        
        # Get file stats
        stat = file_path.stat()
        
        # Create document info
        doc_info = DocumentInfo(
            path=str(file_path),
            filename=file_path.name,
            content=text,
            language=language,
            file_type=file_type,
            size=stat.st_size,
            last_modified=stat.st_mtime
        )
        
        # Cache the document
        self.document_cache[str(file_path)] = doc_info
        
        return doc_info
    
    def scan_documents(self) -> List[DocumentInfo]:
        """Scan and process all documents in the directory"""
        documents = []
        supported_extensions = ['.pdf', '.docx', '.xlsx', '.txt']
        
        logger.info(f"Scanning directory: {self.documents_dir}")
        
        for file_path in self.documents_dir.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                doc_info = self.process_document(file_path)
                if doc_info:
                    documents.append(doc_info)
        
        logger.info(f"Processed {len(documents)} documents")
        return documents
    
    def search_documents(self, query: str, max_results: int = 50) -> List[Dict[str, Any]]:
        """Search for text in documents with improved context"""
        results = []
        query_lower = query.lower()
        
        for doc_info in self.document_cache.values():
            content_lower = doc_info.content.lower()
            
            # Find all occurrences
            start = 0
            matches = []
            while True:
                pos = content_lower.find(query_lower, start)
                if pos == -1:
                    break
                matches.append(pos)
                start = pos + 1
                if len(matches) >= 5:  # Limit matches per document
                    break
            
            if matches:
                for i, pos in enumerate(matches):
                    # Get context around the match
                    context_start = max(0, pos - 150)
                    context_end = min(len(doc_info.content), pos + len(query) + 150)
                    context = doc_info.content[context_start:context_end]
                    
                    # Highlight the match in context
                    context_highlight = context.replace(
                        doc_info.content[pos:pos + len(query)], 
                        f"**{doc_info.content[pos:pos + len(query)]}**"
                    )
                    
                    results.append({
                        "filename": doc_info.filename,
                        "path": doc_info.path,
                        "language": doc_info.language,
                        "file_type": doc_info.file_type,
                        "context": context_highlight,
                        "position": pos,
                        "match_number": i + 1,
                        "total_matches": len(matches),
                        "size": doc_info.size
                    })
                    
                    if len(results) >= max_results:
                        break
            
            if len(results) >= max_results:
                break
        
        return results
    
    def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about the document collection"""
        if not self.document_cache:
            return {"total_documents": 0}
        
        total_size = sum(doc.size for doc in self.document_cache.values())
        languages = {}
        file_types = {}
        
        for doc in self.document_cache.values():
            languages[doc.language] = languages.get(doc.language, 0) + 1
            file_types[doc.file_type] = file_types.get(doc.file_type, 0) + 1
        
        return {
            "total_documents": len(self.document_cache),
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / 1024 / 1024, 2),
            "languages": languages,
            "file_types": file_types,
            "avg_size_kb": round(total_size / len(self.document_cache) / 1024, 2)
        }

# Global document processor - will be initialized in main()
doc_processor = None

# Create the MCP server
server = Server("simple-document-server")

@server.list_tools()
async def handle_list_tools() -> list[types.Tool]:
    """List available tools"""
    return [
        types.Tool(
            name="scan_documents",
            description="Scan and index all documents in the documents directory",
            inputSchema={
                "type": "object",
                "properties": {},
                "required": []
            }
        ),
        types.Tool(
            name="search_documents",
            description="Search for text within the indexed documents",
            inputSchema={
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Text to search for in documents"
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Maximum number of results to return (default: 50)",
                        "default": 50
                    }
                },
                "required": ["query"]
            }
        ),
        types.Tool(
            name="list_documents",
            description="List all processed documents with their metadata",
            inputSchema={
                "type": "object",
                "properties": {},
                "required": []
            }
        ),
        types.Tool(
            name="get_document_stats",
            description="Get statistics about the document collection",
            inputSchema={
                "type": "object",
                "properties": {},
                "required": []
            }
        ),
        types.Tool(
            name="get_document_content",
            description="Get the full content of a specific document",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "Name of the document file"
                    }
                },
                "required": ["filename"]
            }
        )
    ]

@server.call_tool()
async def handle_call_tool(
    name: str, arguments: dict | None
) -> list[types.TextContent]:
    """Handle tool calls"""
    
    try:
        if name == "scan_documents":
            documents = doc_processor.scan_documents()
            result = {
                "status": "success",
                "message": f"Scanned and processed {len(documents)} documents",
                "documents": [
                    {
                        "filename": doc.filename,
                        "path": doc.path,
                        "language": doc.language,
                        "file_type": doc.file_type,
                        "size": doc.size,
                        "content_preview": doc.content[:200] + "..." if len(doc.content) > 200 else doc.content
                    }
                    for doc in documents
                ]
            }
            return [types.TextContent(type="text", text=json.dumps(result, indent=2, ensure_ascii=False))]
        
        elif name == "search_documents":
            if not arguments or "query" not in arguments:
                return [types.TextContent(type="text", text='{"error": "query parameter is required"}')]
            
            query = arguments["query"]
            max_results = arguments.get("max_results", 50)
            results = doc_processor.search_documents(query, max_results)
            
            result = {
                "status": "success",
                "query": query,
                "results_count": len(results),
                "results": results
            }
            
            return [types.TextContent(type="text", text=json.dumps(result, indent=2, ensure_ascii=False))]
        
        elif name == "list_documents":
            documents = list(doc_processor.document_cache.values())
            result = {
                "status": "success",
                "total_documents": len(documents),
                "documents": [
                    {
                        "filename": doc.filename,
                        "path": doc.path,
                        "language": doc.language,
                        "file_type": doc.file_type,
                        "size": doc.size,
                        "content_length": len(doc.content)
                    }
                    for doc in documents
                ]
            }
            return [types.TextContent(type="text", text=json.dumps(result, indent=2, ensure_ascii=False))]
        
        elif name == "get_document_stats":
            stats = doc_processor.get_document_stats()
            result = {
                "status": "success",
                "stats": stats
            }
            return [types.TextContent(type="text", text=json.dumps(result, indent=2, ensure_ascii=False))]
        
        elif name == "get_document_content":
            if not arguments or "filename" not in arguments:
                return [types.TextContent(type="text", text='{"error": "filename parameter is required"}')]
            
            filename = arguments["filename"]
            
            # Find document by filename
            document = None
            for doc in doc_processor.document_cache.values():
                if doc.filename == filename:
                    document = doc
                    break
            
            if not document:
                result = {
                    "status": "error",
                    "message": f"Document '{filename}' not found"
                }
            else:
                result = {
                    "status": "success",
                    "filename": document.filename,
                    "path": document.path,
                    "language": document.language,
                    "file_type": document.file_type,
                    "size": document.size,
                    "content": document.content
                }
            
            return [types.TextContent(type="text", text=json.dumps(result, indent=2, ensure_ascii=False))]
        
        else:
            return [types.TextContent(type="text", text=f'{{"error": "Unknown tool: {name}"}}')]
    
    except Exception as e:
        logger.error(f"Error in tool {name}: {e}")
        return [types.TextContent(type="text", text=f'{{"error": "Internal server error: {str(e)}"}}')]

def parse_arguments():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(
        description="Simple Document MCP Server",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python simple_mcp_server.py                    # Use default ./documents directory
  python simple_mcp_server.py --dir /path/docs   # Use custom directory
  python simple_mcp_server.py -d ~/Documents     # Use home Documents folder
        """.strip()
    )
    
    parser.add_argument(
        "--dir", "-d",
        type=str,
        default="./documents",
        help="Directory containing documents to process (default: ./documents)"
    )
    
    parser.add_argument(
        "--log-level",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        default="INFO",
        help="Set logging level (default: INFO)"
    )
    
    parser.add_argument(
        "--version",
        action="version",
        version="Simple Document MCP Server 1.0.0"
    )
    
    return parser.parse_args()

async def main():
    """Main entry point"""
    global doc_processor
    
    # Parse command line arguments
    args = parse_arguments()
    
    # Set logging level
    logging.getLogger().setLevel(getattr(logging, args.log_level))
    
    # Initialize document processor with specified directory
    doc_processor = SimpleDocumentProcessor(documents_dir=args.dir)
    
    # Validate directory
    if not doc_processor.documents_dir.exists():
        logger.error(f"Documents directory does not exist: {doc_processor.documents_dir}")
        logger.info(f"Creating directory: {doc_processor.documents_dir}")
        doc_processor.documents_dir.mkdir(parents=True, exist_ok=True)
    
    # Initialize and scan documents on startup
    logger.info("Starting Simple Document MCP Server...")
    logger.info(f"Documents directory: {doc_processor.documents_dir.absolute()}")
    logger.info(f"Log level: {args.log_level}")
    
    # Initial scan
    document_count = len(doc_processor.scan_documents())
    logger.info(f"Ready to serve! Found {document_count} documents.")
    
    # Run the server
    async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="simple-document-server",
                server_version="1.0.0",
                capabilities=server.get_capabilities(
                    notification_options=NotificationOptions(),
                    experimental_capabilities={},
                ),
            ),
        )

if __name__ == "__main__":
    asyncio.run(main())